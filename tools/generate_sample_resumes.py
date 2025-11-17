from docx import Document
from fpdf import FPDF
import os

os.makedirs('sample_resumes', exist_ok=True)

# Create DOCX resume 1
doc = Document()
doc.add_heading('Alice Johnson', level=1)
doc.add_paragraph('Software Engineer with 5 years of experience in Python, Django, and REST API development.')
doc.add_paragraph('Skills: Python, Django, REST API, PostgreSQL, Docker, AWS, Git')
doc.save('sample_resumes/alice_johnson.docx')

# Create DOCX resume 2
doc2 = Document()
doc2.add_heading('Bob Smith', level=1)
doc2.add_paragraph('Data Scientist experienced in machine learning, NLP and data analysis.')
doc2.add_paragraph('Skills: Python, Pandas, NumPy, Scikit-learn, TensorFlow, NLP, SQL, Git')
doc2.save('sample_resumes/bob_smith.docx')

# Create PDF resume using FPDF
pdf = FPDF()
pdf.add_page()
pdf.set_font('Arial', 'B', 16)
pdf.cell(0, 10, 'Carol Williams', ln=True)
pdf.set_font('Arial', '', 12)
pdf.multi_cell(0, 8, 'Full Stack Developer with experience in JavaScript, React, Node.js and DevOps practices.')
pdf.ln(2)
pdf.multi_cell(0, 8, 'Skills: JavaScript, React, Node.js, Express, MongoDB, Docker, Kubernetes, AWS, CI/CD, Git')
pdf.output('sample_resumes/carol_williams.pdf')

print('Sample resumes created in sample_resumes/')